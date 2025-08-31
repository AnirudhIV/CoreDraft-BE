from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import models, schemas
from app.database.db import get_db
from typing import Optional, List
from fastapi import Query
from fastapi import Body
from app.database.schemas import GeneratePromptRequest,PromptInput, ContentInput,QuestionInput
from app.utils.ai_generator import generate_document_from_prompt, suggest_tags_ai, summarize_doc_ai,generate_answer_from_context,process_question_and_docs
from app.database.models import Document as DBDocument 
import json
from app.chroma.vectorstore import retrieve_relevant_chunks,add_doc_to_vectorstore,add_text_to_chroma,get_chroma_collection,update_document_embedding,query_similar_docs
from sqlalchemy import func
from langchain.schema import Document as LCDocument
from docx import Document as DocxDocument
import fitz  
from fastapi import APIRouter, UploadFile, File, HTTPException
import os
from app.utils.parser import parse_pdf, parse_docx
from uuid import uuid4
from app.utils.helper import chunk_text, hash_file_content
from app.auth.auth import get_current_user
from chromadb import PersistentClient
from app.database.models import User
from app.utils.text_splitter import hybrid_split_text
from sqlalchemy.exc import SQLAlchemyError
from app.chroma.embedder import embed_text
import traceback
from fastapi import status




import google.generativeai as genai




router = APIRouter()
def admin_required(user:User = Depends(get_current_user)):
    if not user or user.role != "admin":  # adjust to your schema
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Admin privileges required"
        )
    return user

    # Create document
@router.post("/documents", response_model=schemas.DocumentOut)
def upload_document(
    doc: schemas.DocumentCreate,
    db: Session = Depends(get_db),
    chunk_size: int = 500,
    chunk_overlap: int = 100  # Optional
):
    new_doc = models.Document(
        title=doc.title,
        type=doc.type,
        content=doc.content,
        tags=doc.tags,
        vectorstore_status="pending"
    )

    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    try:
        add_doc_to_vectorstore(
            doc_id=str(new_doc.id),
            content=new_doc.content,
            metadata={
                "doc_id": new_doc.id,
                "title": new_doc.title,
                "type": new_doc.type,
            },
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        new_doc.vectorstore_status = "success"
    except Exception as e:
        print(f"❌ Failed to add to vectorstore: {e}")
        new_doc.vectorstore_status = "failed"
    finally:
        db.commit()

    return new_doc



from fastapi import Query
@router.get("/tags", response_model=List[str])
def get_all_tags(db: Session = Depends(get_db)):
    result = db.query(func.jsonb_array_elements_text(models.Document.tags)).distinct().all()
    return sorted({tag for (tag,) in result})


@router.get("/filter", response_model=List[schemas.DocumentOut])
def filter_documents(
    type: Optional[str] = Query(None),
    tag: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    query = db.query(models.Document)

    if type:
        query = query.filter(models.Document.type == type)

    if tag:
        tag_list = [t.strip() for t in tag.split(",")]
        for t in tag_list:
            query = query.filter(models.Document.tags.contains([t]))

    return query.all()


# List all documents
@router.get("/documents", response_model=list[schemas.DocumentOut])
def list_documents(
    tags: Optional[List[str]] = Query(None),
    type: Optional[str] = None,
    source: Optional[str] = None,
    current_user: User = Depends(get_current_user),  # current logged-in user
    db: Session = Depends(get_db),
):
    query = db.query(DBDocument)

    # Only show documents uploaded by the current user
    query = query.filter(DBDocument.user_id == current_user.id)


    # Optional filters
    if tags:
        query = query.filter(DBDocument.tags.overlap(tags))  # PostgreSQL array overlap

    if type:
        query = query.filter(
            (DBDocument.type == type) | (DBDocument.type == "uploaded")
        )

    if source:
        query = query.filter(DBDocument.source == source)

    return query.all()

# Get document by ID
# Get document by ID
@router.get("/documents/{doc_id}", response_model=schemas.DocumentOut)
def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),  # Require auth
    db: Session = Depends(get_db)
):
    document = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Non-admins can't view default docs
    if current_user.role != "admin" and document.is_default:
        raise HTTPException(status_code=403, detail="Not authorized to view this document")

    return document


# Update document

@router.put("/documents/{doc_id}", response_model=schemas.DocumentOut)
def update_document(
    doc_id: int,
    doc: schemas.DocumentUpdate,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Fetch document
    document = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # --- Permission checks ---
    if document.is_default:
        if current_user.role != "admin":
            raise HTTPException(
                status_code=403, detail="Only admins can update default documents"
            )
    else:
        if current_user.role != "admin" and document.owner_id != current_user.id:
            raise HTTPException(
                status_code=403, detail="Not authorized to update this document"
            )

    # --- Track if content changes ---
    content_changed = "content" in doc.dict(exclude_unset=True)

    # --- Apply updates in Postgres ---
    update_data = doc.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(document, key, value)

    db.commit()
    db.refresh(document)

    # --- Sync Chroma if content updated ---
    if content_changed:
        try:
            update_document_embedding(document.id, document.content)
        except Exception as e:
            # Don’t break API if Chroma fails, just warn
            print(f"⚠️ Failed to update Chroma for doc {document.id}: {e}")

    return document



@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    try:
        # Fetch document from Postgres
        document = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete all chunks from ChromaDB
        try:
            client = PersistentClient(path="chroma_db")  # fresh client
            collection = client.get_collection("compliance_docs")  # fresh collection

            # Delete by metadata field 'doc_id' (ensure all chunks have this metadata)
            deleted_count = collection.delete(where={"doc_id": str(doc_id)})
            print(f"✅ Deleted {deleted_count} chunks from ChromaDB for doc_id {doc_id}")

        except Exception as e:
            print("Error deleting vectors from ChromaDB:")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Failed to delete from ChromaDB: {e}")

        # Delete document from Postgres
        try:
            db.delete(document)
            db.commit()
            print(f"✅ Deleted document {doc_id} from Postgres")
        except SQLAlchemyError as e:
            db.rollback()
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

        return {"detail": f"Document {doc_id} deleted successfully from DB and ChromaDB"}

    except HTTPException:
        raise

    except Exception as e:
        db.rollback()
        print("Unexpected error during document deletion:")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Internal Server Error")


# Configure Gemini


 # ✅ adjust the import if needed
@router.post("/generate")
def generate_compliance_doc(
    payload: GeneratePromptRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Step 1: Generate the document using AI
        generated = generate_document_from_prompt(payload.prompt)

        document_id = None

        # Step 2: Save only if explicitly requested
        if getattr(payload, "save", False):
            new_doc = models.Document(
                title=generated["title"],
                content=generated["content"],
                tags=generated["tags"],
                type=payload.type,
                user_id=current_user.id
            )
            db.add(new_doc)
            db.commit()
            db.refresh(new_doc)
            document_id = new_doc.id

            # Add to vectorstore
            try:
                add_doc_to_vectorstore(
                    doc_id=str(new_doc.id),
                    content=new_doc.content,
                    metadata={
                        "title": new_doc.title,
                        "source": "ai",
                        "doc_type": payload.type,
                        "user_id": str(current_user.id),
                    }
                )
            except Exception as e:
                print(f"[Vectorstore] Failed to embed AI-generated doc: {e}")

        # ✅ Step 3: Always return the generated doc (with or without DB id)
        return {
            "status": "success",
            "document_id": document_id,   # null if not saved
            "saved": bool(document_id),
            "document": {
                "title": generated["title"],
                "content": generated["content"],
                "tags": generated["tags"]
            }
        }

    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="AI returned invalid JSON")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))




@router.post("/ai/generate")
def generate_doc(prompt_input: PromptInput, current_user: User = Depends(get_current_user)):
    return generate_document_from_prompt(prompt_input.prompt)

@router.post("/ai/suggest-tags")
def suggest_tags(
    content_input: ContentInput,
    current_user: User = Depends(get_current_user)
):
    return suggest_tags_ai(content_input.content)



@router.post("/ai/summarize")
def summarize_doc(
    content_input: ContentInput,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Get text from DB if doc_id is provided
    if content_input.doc_id:
        doc = (
            db.query(models.Document)
            .filter(
                models.Document.id == content_input.doc_id,
                models.Document.user_id == current_user.id
            )
            .first()
        )
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        text = doc.content

    # Or use provided content
    elif content_input.content:
        text = content_input.content

    else:
        raise HTTPException(status_code=400, detail="Either doc_id or content must be provided")

    # Generate summary using AI
    summary = summarize_doc_ai(text)
    return {"summary": summary}

@router.post("/ask")
def ask_question(request: QuestionInput, current_user: User = Depends(get_current_user)):
    question = request.question.strip()
    if not question:
        return {"answer": "Please enter a valid question.", "sources": []}

    top_k = 5  # number of top chunks to retrieve

    try:
        answer, sources = process_question_and_docs(question, top_k=top_k)
    except Exception as e:
        return {"answer": f"An error occurred while generating the answer: {str(e)}", "sources": []}

    return {"answer": answer, "sources": sources}



UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)  # ✅ Ensure upload directory exists




@router.post("/upload", tags=["Upload"])
async def upload_doc(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        MAX_USER_DOCS = 3 # Maximum number of user-uploaded documents allowed

        # Count how many non-default docs uploaded by this user
        user_doc_count = db.query(DBDocument).filter(
            DBDocument.user_id == current_user.id,
            DBDocument.is_default == False
        ).count()

        if user_doc_count >= MAX_USER_DOCS:
            raise HTTPException(
                status_code=400,
                detail=f"Upload limit reached: Maximum {MAX_USER_DOCS} uploaded documents allowed."
            )

        # Read & hash file
        file_bytes = await file.read()
        file_hash = hash_file_content(file_bytes)
        file_path = os.path.join(UPLOAD_DIR, file.filename)

        with open(file_path, "wb") as f:
            f.write(file_bytes)

        # Extract text
        if file.filename.lower().endswith(".pdf"):
            content = parse_pdf(file_path)
        elif file.filename.lower().endswith(".docx"):
            content = parse_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        if not content.strip():
            raise ValueError("Parsed text is empty")

        # AI tag suggestion first
        ai_tags = suggest_tags_ai(content)

        # Save to DB first so we have an ID for Chroma metadata
        new_doc = DBDocument(
            title=file.filename,
            content=content,
            type="uploaded",
            tags=ai_tags,
            user_id=current_user.id,
            is_default=(current_user.role == "admin")
        )

        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)  # Now new_doc.id is available

        # Chunk & Embed (store doc_id + tags in metadata)
        chunks = hybrid_split_text(content)
        for i, chunk in enumerate(chunks):
            add_text_to_chroma(
                text=chunk["text"],
                doc_id=f"{new_doc.id}-{i}",  # unique chunk id
                metadata={
                    "source": file.filename,
                    "file_hash": file_hash,
                    "user_id": current_user.id,
                    "doc_id": str(new_doc.id),  # for filtering/deleting later
                    "tags":  ", ".join(ai_tags),
                    **chunk.get("metadata", {})
                }
            )

        print(f"✅ Upload succeeded: {file.filename} | Chunks: {len(chunks)}")

        return {
            "message": "File uploaded, AI tags added, and document saved.",
            "chunks_added": len(chunks),
            "doc_id": new_doc.id,
            "tags": ai_tags,
            "title": new_doc.title,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))







@router.delete("/rag/reset", tags=["RAG"])
def reset_rag_data(user=Depends(admin_required)):  # <-- only admins allowed
    client = PersistentClient(path="chroma_db")  # Adjust if different
    collections = client.list_collections()
    for col in collections:
        client.delete_collection(name=col.name)
    return {"message": "All RAG collections deleted."}

@router.get("/rag/status", tags=["RAG"])
def rag_status():
    """
    Returns the current number of vectors in each Chroma collection.
    Always uses a fresh PersistentClient and collection instance to avoid stale data.
    """
    client = PersistentClient(path="chroma_db")  # fresh client each call
    status = {}

    for col_obj in client.list_collections():
        try:
            # Get a fresh collection instance
            collection = client.get_collection(col_obj.name)
            count = collection.count()
            status[col_obj.name] = count
        except Exception as e:
            # In case fetching the collection fails
            status[col_obj.name] = f"Error: {str(e)}"

    return {"rag_collections": status}


@router.get("/debug/collections")
def list_collections():
    client = PersistentClient(path="chroma_db")
    collections = [col.name for col in client.list_collections()]
    return {"collections": collections}


